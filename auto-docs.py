import configparser
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm


# хранение классов из исходника
class class_:
    class_name = ''
    fields = []
    methods = []

    def __init__(self, class_name, fields, methods):
        self.class_name = class_name
        self.fields = fields
        self.methods = methods

    def get_all(self):
        return {'class_name': self.class_name, 'fields': self.fields, 'methods': self.methods}


lab_names = [
    "Знакомство с интегрированной средой Code::Blocks",
    "Организация циклов в С++",
    "Функции. Передача параметра по значению. Перегрузка",
    "Функции. Передача параметра по ссылке",
    "Обработка числовых последовательностей с использованием текстовых файлов",
    "Списки, массивы, векторы",
    "Матрицы",
    "Строки. Процедурная и объектно-ориентированная библиотеки. Широкие строки.",
    "Бинарные файлы",
    "Структуры и файлы",
    "Многомодульные проекты",
    "Обработка строк, хранящихся в файле",
    "Множества",
    "Рекурсивные функции",
    "Обработка исключений. Функции, генерирующие исключения",
    "Указатели на функции",
    "Реализация стека и очереди на базе списка, на базе вектора и на базе динамического массива",
    "Длинная арифметика",
    "Динамические структуры данных и файлы",
    "Структура-пара",
    "Функции с переменным числом параметров",
    "Объединения. Перечисления. Типы, определяемые пользователем",
    "Поразрядные операции и битовые поля",
    "Шаблоны функций"
]

purposes = [
    "Получение первичных навыков работы в интегрированной среде Code::Blocks",
    "выработка навыков программирования циклических вычислительных процессов;\nпорядок работы с вложенными циклами;\nформатированный вывод данных.",
    "закрепление навыков в организации итерационных и арифметических циклов, использования вспомогательных алгоритмов (функций с передачей параметров по значению).",
    "выработка навыков передачи параметров в функцию по ссылке.",
    "\nприобретение навыков работы с текстовыми файлами;\nзакрепление навыков в организации итерационных и арифметических циклов, использования вспомогательных алгоритмов (функций с передачей параметров по значению). ",
    "Знакомство с динамическими информационными структурами на примере одно- и двунаправленных списков, динамических массивов и векторов",
    "выработка навыков реализации матриц в языке C++ различными способами и закрепление навыков реализации типовых алгоритмов (классификация, поиск, сортировка).",
    "Получить практические навыки работы с библиотеками cstring, string, а также работы с широкими строками и символами wstring.",
    "выработка навыков работы с бинарными файлами",
    "1. Получить практические навыки работы со структурами и файлами.\n2. Получить практические навыки организации массивов/векторов с элементами сложной структуры.",
    None,
    None,
    "Изучить правила описания и использования контейнера Set из стандартной библиотеки шаблонов. Получить навыки в решении практических задач и ребусов с использованием теории множеств.",
    "Получить практические навыки работы с рекурсивными функциями.",
    None,
    None,
    "1.Изучить алгоритмы формирования и просмотра динамической структуры данных – «стек». Научиться программировать различные действия над его элементами\n2.Изучить алгоритмы формирования и просмотра динамической структуры данных «очередь». Научиться программировать различные действия над ее элементами.",
    None,
    None,
    "закрепление навыков использования структур.",
    "Приобрести практические навыки работы с функциями с переменным числом параметров и закрепить использование указателей на функции.",
    "изучить синтаксис и правила работы с объединениями, перечислениями и типами, определяемые пользователем",
    "изучить теорию и научиться программировать поразрядные операции и битовые поля.",
    "научиться создавать шаблоны функций для работы с любыми типами данных без переписывания кода программы."
]


# создание шабки таблицы функций
def set_func_table(tab):
    row_cells = tab.row_cells(0)
    tab.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tab.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    func = row_cells[0].paragraphs[0].add_run('Функция')
    desc = row_cells[1].paragraphs[0].add_run('Назначение')
    func.bold = True
    desc.bold = True


# создание шабки таблицы полей
def set_fields_table(tab):
    row_cells = tab.row_cells(0)
    for i in range(3):
        tab.cell(0, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    name = row_cells[0].paragraphs[0].add_run('Имя')
    type = row_cells[1].paragraphs[0].add_run('Тип')
    desc = row_cells[2].paragraphs[0].add_run('Назначение')
    name.bold = True
    type.bold = True
    desc.bold = True


# создание шабки таблицы методов
def set_methods_table(tab):
    row_cells = tab.row_cells(0)
    tab.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tab.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    func = row_cells[0].paragraphs[0].add_run('Метод')
    desc = row_cells[1].paragraphs[0].add_run('Назначение')
    func.bold = True
    desc.bold = True


# поиск и сохранение всех функций в исходнике
def get_funcs_from_cpp(path):
    funcs_list = []
    with open(path) as src:
        for line in src:
            if line.startswith('/*func'):
                description = line
                description = description.rstrip()
                description = description.replace('/*func ', '')
                description = description.replace('*/', '')
                if '*/' not in line:
                    for line in src:
                        if '*/' in line:
                            break
                code = next(src).rstrip()
                funcs_list.append((code, description))
    return funcs_list


# поиск и сохранение всех классов в исходнике
def get_classes_from_cpp(path):
    classes_list = []
    with open(path) as src:
        for line in src:
            if 'class' in line:
                class_name = line.split(' ')[1]
                class_obj = class_(class_name, [], [])
                while '};' not in line:
                    if '//f' in line:
                        line = line.strip().split(' ')
                        f_type = line[0]
                        f_name = line[1].replace(';', '')
                        f_desc = ' '.join([line[i] for i in range(4, len(line))])
                        field = {'field_name': f_name, 'data_type': f_type, 'description': f_desc}
                        class_obj.fields.append(field.copy())
                    if '/*m' in line:
                        line = line.strip().replace('/*m ', '').replace('*/', '')
                        m_desc = line
                        m_func = next(src).strip()
                        method = {'func': m_func, 'description': m_desc}
                        class_obj.methods.append(method.copy())
                    line = next(src)
                classes_list.append(class_obj)
    return classes_list


# заполнение таблицы функций
def fill_funcs_table(tab, funcs):
    for i in range(len(funcs)):
        new_row = tab.add_row()
        new_row.cells[0].paragraphs[0].text = funcs[i][0]
        new_row.cells[1].paragraphs[0].text = funcs[i][1]


# заполнение таблицы полей
def fill_fields_table(tab, fields):
    for i in range(len(fields)):
        new_row = tab.add_row()
        new_row.cells[0].paragraphs[0].text = fields[i]['field_name']
        new_row.cells[1].paragraphs[0].text = fields[i]['data_type']
        new_row.cells[2].paragraphs[0].text = fields[i]['description']


# заполнение таблицы методов
def fill_methods_table(tab, methods):
    for i in range(len(methods)):
        new_row = tab.add_row()
        new_row.cells[0].paragraphs[0].text = methods[i]['func']
        new_row.cells[1].paragraphs[0].text = methods[i]['description']


# установка стиля для таблицы
def set_table_style(tab):
    # tab.allow_autofit = False
    for row in tab.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            # cell.width = Cm(8.74)
            for paragraph in paragraphs:
                paragraph.paragraph_format.line_spacing = 1.5
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    font = run.font
                    font.size = Pt(12)


# установка нужного стиля для абзаца
def set_p_style(p, font_size, is_bold):
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = 0
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.bold = is_bold
        font = run.font
        font.size = Pt(font_size)


# создание ini файла
def create_config(path):
    config = configparser.ConfigParser()
    config.add_section("Параметры")
    config.set("Параметры", "Номер лабы", "")
    config.set("Параметры", "Номер варианта", "")
    config.set("Параметры", "Автор", "")
    config.set("Параметры", "Преподаватель", "")
    config.set("Параметры", "Имя исходника", "")

    with open(path, "w") as config_file:
        config.write(config_file)


# чтение параметров из ini файла
def read_config(path):
    config = configparser.ConfigParser()
    config.read(path)
    cfg = []
    cfg.append(config.get("Параметры", "Номер лабы"))
    cfg.append(config.get("Параметры", "Номер варианта"))
    cfg.append(config.get("Параметры", "Автор"))
    cfg.append(config.get("Параметры", "Преподаватель"))
    cfg.append(config.get("Параметры", "Имя исходника"))
    return cfg


path = "config.ini"
if not os.path.exists(path):
    create_config(path)
    print("config.ini создан.")
    exit(0)

cfg = read_config(path)
lab_number = cfg[0]
if lab_number != '':
    lab_number = int(lab_number)
var = cfg[1]
if var != '':
    var = int(var)
author_name = cfg[2]
teacher_name = cfg[3]
cpp_path = cfg[4]

output_filename = cpp_path.split(".")[0] + '.docx'

funcs = get_funcs_from_cpp(cpp_path)
get_classes_from_cpp(cpp_path)
classes = get_classes_from_cpp(cpp_path)
document = Document()
document._body.clear_content()

if lab_number != '' and var != '':
    header_ = document.add_paragraph()
    header_.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = header_.add_run()
    r.font.size = Pt(18)
    r.bold = True
    r.font.name = 'Times New Roman'

    r2 = header_.add_run()
    r2.font.size = Pt(14)
    r2.bold = True
    r2.font.name = 'Times New Roman'

    r.text = 'Отчёт\xa0по\xa0лабораторной\xa0работе\xa0№\xa0{lab} Вариант\xa0{var}\n'.format(lab=lab_number, var=var)
    r2.text = 'Дисциплина\xa0«Программирование\xa0и\xa0информатика», 2\xa0семестр\n«' + lab_names[lab_number - 1] + '»'

    author = document.add_paragraph()
    author.paragraph_format.line_spacing = 1.5
    author_r1 = author.add_run()
    author_r1.text = 'Выполнил \n'
    author_r1.bold = True
    author_r1.font.size = Pt(12)
    author_r1.font.name = 'Times New Roman'

    author_r2 = author.add_run()
    author_r2.text = 'студент группы ДИПРб11 ____________ {author_name} «____»__________ 2020\n'.format(
        author_name=author_name)
    author_r2.font.name = 'Times New Roman'
    author_r2.font.size = Pt(12)

    author_r3 = author.add_run()
    author_r3.text = 'Проверила \n'
    author_r3.bold = True
    author_r3.font.name = 'Times New Roman'
    author_r3.font.size = Pt(12)

    author_r4 = author.add_run()
    author_r4.text = 'ст. преп. кафедры АСОИУ ____________ {teacher_name} «____»__________ 2020'.format(
        teacher_name=teacher_name)
    author_r4.font.name = 'Times New Roman'
    author_r4.font.size = Pt(12)

    if purposes[lab_number - 1] is not None:
        purpose = document.add_paragraph()
        purpose_r1 = purpose.add_run()
        purpose_r1.text = "Цель работы: "
        purpose_r1.font.name = 'Times New Roman'
        purpose_r1.font.size = Pt(12)
        purpose_r1.bold = True

        purpose_r2 = purpose.add_run()
        purpose_r2.text = purposes[lab_number - 1]
        purpose_r2.font.name = 'Times New Roman'
        purpose_r2.font.size = Pt(12)

    document.add_paragraph()

if len(funcs) > 0:
    func_label = document.add_paragraph('Таблица X.X - Функции, обеспечивающие работу программы')
    set_p_style(func_label, font_size=12, is_bold=False)

    func_tab = document.add_table(1, 2, 'TableGrid')
    set_func_table(func_tab)
    fill_funcs_table(func_tab, funcs)
    set_table_style(func_tab)
    document.add_paragraph()

if len(classes) > 0:
    for i in range(len(classes)):
        class_label = 'Таблица X.X - Описание класса {class_name}'.format(class_name=classes[i].get_all()['class_name'])
        class_label = document.add_paragraph(class_label)
        set_p_style(class_label, 12, False)
        fields_table = document.add_table(1, 3, 'TableGrid')
        set_fields_table(fields_table)
        fields = classes[i].get_all()['fields']
        methods = classes[i].get_all()['methods']
        fill_fields_table(fields_table, fields)

        method_table = document.add_table(1, 2, 'TableGrid')
        set_methods_table(method_table)
        fill_methods_table(method_table, methods)

        set_table_style(fields_table)
        set_table_style(method_table)

        document.add_paragraph()


document.save(output_filename)
